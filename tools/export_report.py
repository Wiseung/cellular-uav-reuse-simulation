from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "完整中文报告初稿.md"
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_DOCX = OUTPUT_DIR / "wireless_reuse_report_final.docx"
OUTPUT_PDF = OUTPUT_DIR / "wireless_reuse_report_final.pdf"
OUTPUT_SUBMISSION_DOCX = OUTPUT_DIR / "wireless_reuse_report_submission.docx"
OUTPUT_SUBMISSION_PDF = OUTPUT_DIR / "wireless_reuse_report_submission.pdf"
COURSE_NAME = "现代通信的无线传播信道"
ASSIGNMENT_TYPE = "课程大作业（仿真类）"
SUBMISSION_DATE = "2026年7月8日"


def set_run_font(run, size_pt: float | None = None, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.different_first_page_header_footer = True

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(11)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    heading1.font.size = Pt(12.5)
    heading1.font.bold = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    heading2.font.size = Pt(11.5)
    heading2.font.bold = True

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_paragraph)


def convert_latex_to_text(text: str) -> str:
    output = text
    replacements = {
        r"\quad": " ",
        r"\,": " ",
        r"\%": "%",
        r"\ge": "≥",
        r"\le": "≤",
        r"\approx": "≈",
        r"\propto": "∝",
        r"\times": "×",
        r"\in": "∈",
        r"\to": "→",
    }
    for source, target in replacements.items():
        output = output.replace(source, target)

    output = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", output)
    output = re.sub(r"\\mathbb\{([^{}]+)\}", r"\1", output)
    output = re.sub(r"\\mathbf\{([^{}]+)\}", r"\1", output)
    output = re.sub(r"\\text\{([^{}]+)\}", r"\1", output)

    while True:
        updated = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", output)
        if updated == output:
            break
        output = updated
    while True:
        updated = re.sub(r"\\sqrt\{([^{}]+)\}", r"sqrt(\1)", output)
        if updated == output:
            break
        output = updated

    output = re.sub(r"_\{([^{}]+)\}", r"_\1", output)
    output = re.sub(r"\^\{([^{}]+)\}", r"^\1", output)
    output = output.replace("{", "").replace("}", "")
    output = output.replace("\\", "")
    return output.strip()


def normalize_inline_math(text: str) -> str:
    def _replace(match: re.Match[str]) -> str:
        return convert_latex_to_text(match.group(1))

    return re.sub(r"\$([^$]+)\$", _replace, text)


def clean_text(text: str) -> str:
    return normalize_inline_math(text).strip()


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, 10)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text)
    run._r.append(fld_char_end)


def add_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, size=11, first_line_indent_cm: float | None = 0.74) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if first_line_indent_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run = paragraph.add_run(clean_text(text))
    set_run_font(run, size, bold=bold, italic=italic)


def add_label_paragraph(doc: Document, label: str, body: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    bold_run = paragraph.add_run(label)
    set_run_font(bold_run, 10.5, bold=True)
    text_run = paragraph.add_run(clean_text(body))
    set_run_font(text_run, 10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.strip())
    set_run_font(run, 12.5 if level == 1 else 11.5, bold=True)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_equation(doc: Document, equation_lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    equation_text = " ".join(line.strip() for line in equation_lines if line.strip() not in {r"\[", r"\]"})
    run = paragraph.add_run(convert_latex_to_text(equation_text))
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(11.5)


def parse_markdown_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start_index
    while index < len(lines):
        line = lines[index].strip()
        if not line.startswith("|"):
            break
        cells = [clean_text(cell.strip()) for cell in line.strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r"[:\- ]+", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(8)
    caption_paragraph.paragraph_format.space_after = Pt(4)
    caption_text = re.sub(r"\*\*", "", caption).strip()
    caption_text = re.sub(r"^(TABLE\s+[IVXLC]+)\s+", r"\1. ", caption_text)
    caption_run = caption_paragraph.add_run(caption_text)
    set_run_font(caption_run, 10.5, bold=True)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, 10, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption.strip("*"))
    set_run_font(caption_run, 10, italic=True)


def add_reference(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.first_line_indent = Cm(-0.74)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(clean_text(text))
    set_run_font(run, 10.5)


def add_cover_page(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(70)
    run = paragraph.add_run(COURSE_NAME)
    set_run_font(run, 18, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(ASSIGNMENT_TYPE)
    set_run_font(run, 14, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(36)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    set_run_font(run, 20, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(30)
    run = paragraph.add_run(subtitle)
    set_run_font(run, 13, italic=True)

    fields = [
        ("学生姓名", "________________"),
        ("学号", "________________"),
        ("班级", "________________"),
        ("任课教师", "________________"),
        ("提交日期", SUBMISSION_DATE),
    ]
    for label, value in fields:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(f"{label}：{value}")
        set_run_font(run, 13)

    add_page_break(doc)


def add_toc_page(doc: Document, entries: list[tuple[str, int]]) -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(20)
    heading.paragraph_format.space_after = Pt(18)
    run = heading.add_run("目录")
    set_run_font(run, 16, bold=True)

    for title, page in entries:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(f"{title} ................................ {page}")
        set_run_font(run, 11)
    add_page_break(doc)


def build_docx(source_path: Path, output_docx: Path, *, include_front_matter: bool, toc_entries: list[tuple[str, int]] | None = None) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    if include_front_matter:
        title_line = next((line[2:].strip() for line in lines if line.startswith("# ")), "课程作业报告")
        subtitle_line = next((line.strip("*").strip() for line in lines if line.startswith("*") and line.endswith("*")), "")
        add_cover_page(doc, title_line, subtitle_line)
        add_toc_page(doc, toc_entries or [])

    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue

        if raw_line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(raw_line[2:].strip())
            set_run_font(run, 16, bold=True)
            index += 1
            continue

        if line.startswith("*") and line.endswith("*") and not line.startswith("*Fig."):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(10)
            run = paragraph.add_run(clean_text(line.strip("*")))
            set_run_font(run, 11.5, italic=True)
            index += 1
            continue

        if line.startswith("**摘要—**"):
            add_label_paragraph(doc, "摘要—", line.replace("**摘要—**", "", 1).strip())
            index += 1
            continue

        if line.startswith("**关键词—**"):
            add_label_paragraph(doc, "关键词—", line.replace("**关键词—**", "", 1).strip())
            index += 1
            continue

        if raw_line.startswith("## "):
            add_heading(doc, raw_line[3:], level=1)
            index += 1
            continue

        if raw_line.startswith("### "):
            add_heading(doc, raw_line[4:], level=2)
            index += 1
            continue

        if line == r"\[":
            equation_lines = []
            while index < len(lines):
                equation_lines.append(lines[index])
                if lines[index].strip() == r"\]":
                    break
                index += 1
            add_equation(doc, equation_lines)
            index += 1
            continue

        if line.startswith("**TABLE "):
            caption = line.strip("*")
            probe = index + 1
            while probe < len(lines) and not lines[probe].strip():
                probe += 1
            rows, next_index = parse_markdown_table(lines, probe)
            add_table(doc, caption, rows)
            index = next_index
            continue

        if line.startswith("![Fig."):
            image_match = re.match(r"!\[.*?\]\((.+)\)", line)
            if image_match is None:
                raise ValueError(f"Invalid image line: {line}")
            image_path = ROOT / image_match.group(1)
            probe = index + 1
            while probe < len(lines) and not lines[probe].strip():
                probe += 1
            caption = lines[probe].strip() if probe < len(lines) and lines[probe].strip().startswith("*Fig.") else ""
            add_figure(doc, image_path, caption)
            index = probe + 1 if caption else index + 1
            continue

        if re.match(r"\[\d+\]", line):
            add_reference(doc, line)
            index += 1
            continue

        add_paragraph(doc, line)
        index += 1

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    ps_script = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open('{docx_path.resolve()}')
$doc.Fields.Update() | Out-Null
$doc.SaveAs([ref] '{pdf_path.resolve()}', [ref] 17)
$doc.Close()
$word.Quit()
"""
    subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps_script],
        check=True,
        cwd=str(ROOT),
    )


def extract_section_pages_from_pdf(pdf_path: Path) -> list[tuple[str, int]]:
    section_titles = [
        "I. 引言",
        "II. 理论主体与总体研究进展情况",
        "III. 仿真场景与实验设计",
        "IV. 仿真实验结果与真实站址数据分析",
        "V. 分析与讨论",
        "VI. 结论与展望",
        "VII. 参考文献",
        "VIII. 附录",
    ]
    reader = PdfReader(str(pdf_path))
    entries: list[tuple[str, int]] = []
    for title in section_titles:
        found_page = None
        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if title in page_text:
                found_page = page_index
                break
        if found_page is not None:
            entries.append((title, found_page + 2))
    return entries


def main() -> None:
    build_docx(SOURCE_MD, OUTPUT_DOCX, include_front_matter=False)
    convert_docx_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)
    print(f"Wrote DOCX: {OUTPUT_DOCX}")
    print(f"Wrote PDF: {OUTPUT_PDF}")
    toc_entries = extract_section_pages_from_pdf(OUTPUT_PDF)
    build_docx(SOURCE_MD, OUTPUT_SUBMISSION_DOCX, include_front_matter=True, toc_entries=toc_entries)
    convert_docx_to_pdf(OUTPUT_SUBMISSION_DOCX, OUTPUT_SUBMISSION_PDF)
    print(f"Wrote DOCX: {OUTPUT_SUBMISSION_DOCX}")
    print(f"Wrote PDF: {OUTPUT_SUBMISSION_PDF}")


if __name__ == "__main__":
    main()
